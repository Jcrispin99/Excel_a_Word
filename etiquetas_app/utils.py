import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor # <--- Añadir RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime # <--- Añadir esta importación
import qrcode # <--- Añadir esta importación
import io     # <--- Añadir esta importación
from PIL import Image

def set_table_borders(table):
    # Tu código existente para configurar bordes
    tbl = table._element
    tblPr = tbl.xpath("w:tblPr")
    if not tblPr:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    else:
        tblPr = tblPr[0]
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def remove_paragraph_spacing(paragraph):
    # Tu código existente para eliminar espacios
    if paragraph:
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '20')
        spacing.set(qn('w:after'), '20')
        pPr.append(spacing)

def adjust_cell_spacing(cell):
    # Tu código existente para ajustar espacios en celdas
    tcPr = cell._element.tcPr
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        cell._element.append(tcPr)
    
    tblCellMar = OxmlElement('w:tcMar')
    margins = {'top': '40', 'left': '60', 'bottom': '40', 'right': '60'}
    for margin, value in margins.items():
        mar = OxmlElement(f'w:{margin}')
        mar.set(qn('w:w'), value)
        mar.set(qn('w:type'), 'dxa')
        tblCellMar.append(mar)
    
    existing_mar = tcPr.find(qn('w:tcMar'))
    if existing_mar is not None:
        tcPr.remove(existing_mar)
    tcPr.append(tblCellMar)
    
    tcVAlign = OxmlElement('w:vAlign')
    tcVAlign.set(qn('w:val'), 'top')
    tcPr.append(tcVAlign)
    
    for paragraph in cell.paragraphs:
        remove_paragraph_spacing(paragraph)

def generate_word_document(excel_path, images_dir, output_path):
    # Leer el archivo Excel
    df = pd.read_excel(excel_path)
    doc = Document()
    
    # Configurar márgenes del documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Procesar cada fila del Excel
    for index, row in df.iterrows():
        # Buscar la imagen correspondiente en el directorio de imágenes
        imagen_path = None
        if 'IMAGEN' in row and pd.notna(row['IMAGEN']):
            imagen_nombre = str(row['IMAGEN']).strip() # Obtener el valor del Excel y limpiar espacios

            if imagen_nombre: # Solo proceder si imagen_nombre no está vacío después de strip()
                # Si imagen_nombre no contiene un punto (es decir, no tiene extensión), añadir '.jpg'
                if '.' not in imagen_nombre:
                    imagen_nombre += '.PNG'
                
                # Ahora, imagen_nombre contiene el nombre del archivo a buscar (ej: "codigo.jpg" o "imagen.png")
                # Buscar el archivo en el directorio de imágenes (y subdirectorios)
                for root, dirs, files_in_current_dir in os.walk(images_dir):
                    for file_name in files_in_current_dir:
                        # Comparar el nombre del archivo encontrado con imagen_nombre (ignorando mayúsculas/minúsculas)
                        if file_name.lower() == imagen_nombre.lower():
                            imagen_path = os.path.join(root, file_name)
                            break # Imagen encontrada, salir del bucle de archivos
                    if imagen_path:
                        break # Imagen encontrada, salir del bucle de directorios
        
        # Obtener el número de cajas
        num_cajas = int(row.get('CONTEO_CAJAS', 1))
        
        # Generar una tabla para cada caja
        for num_caja in range(1, num_cajas + 1):
            # Crear tabla principal
            table = doc.add_table(rows=2, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.columns[0].width = Inches(3)
            table.columns[1].width = Inches(3)
            set_table_borders(table)
            
            # Ajustar espacios en todas las celdas
            for row_cells in table.rows:
                for cell in row_cells.cells:
                    adjust_cell_spacing(cell)
            
            # Celda izquierda: imagen y datos
            left_cell = table.cell(0, 0)
            
            # Modificar la parte donde se inserta la imagen (alrededor de la línea 120)
            # Agregar imagen si existe
            if imagen_path and os.path.exists(imagen_path):
                try:
                    p = left_cell.paragraphs[0]
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = p.add_run()
                    
                    # Obtener dimensiones originales de la imagen
                    with Image.open(imagen_path) as img:
                        width_orig, height_orig = img.size
                        ratio = height_orig / width_orig
                    
                    # Establecer ancho fijo y calcular altura proporcional
                    width_inches = 2.5  # Ancho fijo en pulgadas
                    
                    # Establecer un límite máximo para la altura (por ejemplo, 3 pulgadas)
                    max_height_inches = 3.0
                    
                    # Calcular altura proporcional
                    height_inches = width_inches * ratio
                    
                    # Si la altura calculada excede el máximo, ajustar proporcionalmente
                    if height_inches > max_height_inches:
                        height_inches = max_height_inches
                        width_inches = height_inches / ratio
                    
                    # Insertar imagen con las dimensiones calculadas
                    run.add_picture(imagen_path, width=Inches(width_inches), height=Inches(height_inches))
                except Exception as e:
                    p = left_cell.paragraphs[0]
                    p.text = f"Error al procesar imagen: {str(e)}"
            
            # Tabla de datos
            # La tabla ahora necesitará más filas para acomodar los merges y nuevos campos
            # CÓDIGO (4) + REVISADO (2) + BLANCO (1) + RECEPCIONADO (1) + DESCRIPCIÓN (1) + CAJA (1) + CANTIDAD (1) = 11 filas
            data_table = left_cell.add_table(rows=11, cols=1)
            data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(data_table)
            
            # Aplicar ajuste de espaciado a todas las celdas de data_table
            for r_idx, row_cells_dt in enumerate(data_table.rows):
                for c_idx, cell_dt in enumerate(row_cells_dt.cells):
                    adjust_cell_spacing(cell_dt)
            
            # 1. Celda para CÓDIGO (ocupa 4 filas)
            codigo_cell = data_table.cell(0, 0)
            
            # Limpiar el contenido por defecto y añadir el nuevo con formato
            codigo_cell.text = '' # Limpiar cualquier texto previo
            
            # Párrafo para el texto del CÓDIGO
            p_codigo = codigo_cell.paragraphs[0] # Usar el párrafo existente o añadir uno nuevo si es necesario
            if not codigo_cell.paragraphs: # Asegurar que haya al menos un párrafo
                 p_codigo = codigo_cell.add_paragraph()
            else:
                 p_codigo = codigo_cell.paragraphs[0]
            
            p_codigo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run_codigo = p_codigo.add_run(str(row['CODIGO']))
            run_codigo.font.size = Pt(33)
            run_codigo.font.color.rgb = RGBColor(255, 0, 0) # Color Rojo

            # Generar y añadir el Código QR debajo del texto del código
            qr_data = str(row['CODIGO'])
            qr_img = qrcode.make(qr_data)
            
            # Guardar QR en un stream en memoria
            qr_image_stream = io.BytesIO()
            qr_img.save(qr_image_stream, format='PNG')
            qr_image_stream.seek(0) # Rebobinar el stream al principio

            # Añadir un nuevo párrafo para el QR en la misma celda
            p_qr = codigo_cell.add_paragraph() 
            p_qr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run_qr = p_qr.add_run()
            # Ajusta el tamaño del QR según necesites, por ejemplo Inches(1.0) o Inches(1.5)
            run_qr.add_picture(qr_image_stream, width=Inches(2.0)) 

            # Fusionar las celdas para CÓDIGO (esto ya lo tenías)
            codigo_cell.merge(data_table.cell(1, 0))
            codigo_cell.merge(data_table.cell(2, 0))
            codigo_cell.merge(data_table.cell(3, 0))
            
            # Ajustar alineación vertical de la celda fusionada de CÓDIGO a 'center'
            tcPr_codigo = codigo_cell._element.get_or_add_tcPr()
            # Eliminar alineación vertical previa si existe (de adjust_cell_spacing)
            existing_vAlign = tcPr_codigo.find(qn('w:vAlign'))
            if existing_vAlign is not None:
                tcPr_codigo.remove(existing_vAlign)
            # Añadir nueva alineación vertical
            vAlign_codigo = OxmlElement('w:vAlign')
            vAlign_codigo.set(qn('w:val'), 'center') # Centrar verticalmente
            tcPr_codigo.append(vAlign_codigo)

            # 2. Celda para REVISADO POR (ocupa 2 filas)
            revisado_cell = data_table.cell(4, 0)
            revisado_cell.text = "REVISADO POR:"
            # Fusionar las celdas para REVISADO POR
            revisado_cell.merge(data_table.cell(5, 0))

            # 3. Celda en BLANCO (ocupa 1 fila)
            # data_table.cell(6, 0).text = "" # Ya está vacía por defecto

            # 4. Celda para RECEPCIONADO (ocupa 1 fila)
            current_date = datetime.now().strftime("%d/%m/%Y")
            data_table.cell(7, 0).text = f"RECEPCIONADO: {current_date}"
            
            # 5. Celda para DESCRIPCIÓN (ocupa 1 fila)
            data_table.cell(8, 0).text = f"DESCRIPCIÓN: {row['DESCRIPCION']}"
            
            # 6. Celda para CAJA (ocupa 1 fila)
            data_table.cell(9, 0).text = f"CAJA {num_caja} DE {num_cajas}"
            
            # 7. Celda para CANTIDAD (ocupa 1 fila)
            data_table.cell(10, 0).text = f"CANTIDAD: {row['CANTIDAD']}"
            
            # Celda derecha: tabla de registro
            right_cell = table.cell(0, 1)
            
            # Eliminar párrafo existente
            if right_cell.paragraphs:
                p = right_cell.paragraphs[0]
                p._element.getparent().remove(p._element)
                p._p = None
            
            # Tabla interna
            inner_table = right_cell.add_table(rows=22, cols=4)
            inner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(inner_table)
            
            for row_cells in inner_table.rows:
                for cell in row_cells.cells:
                    adjust_cell_spacing(cell)
            
            headers = ['FECHA DD/MM/AA', 'CANT.', 'RP.', 'FIRMA']
            for j, header in enumerate(headers):
                inner_table.cell(0, j).text = header
            
            for i in range(1, 22):
                for j in range(4):
                    inner_table.cell(i, j).text = ""
            
            # Celda de observación
            obs_cell = table.cell(1, 0)
            obs_cell.merge(table.cell(1, 1))
            obs_cell.text = "OBSERVACIÓN:"
            
            # Aumentar altura de la celda de observación
            tr = obs_cell._element.getparent()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '800')
            trHeight.set(qn('w:hRule'), 'atLeast')
            trPr = tr.get_or_add_trPr()
            trPr.append(trHeight)
            
            # Agregar salto de página excepto en la última tabla
            if not (index == len(df) - 1 and num_caja == num_cajas):
                doc.add_page_break()
    
    # Guardar el documento
    doc.save(output_path)
    return output_path