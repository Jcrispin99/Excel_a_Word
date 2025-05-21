import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Crear el directorio 'recursos' si no existe
os.makedirs('recursos', exist_ok=True)

# Leer el archivo Excel
df = pd.read_excel('LISTA1.xlsx')

# Crear un nuevo documento Word desde cero
doc = Document()
# Ajustar márgenes del documento (no demasiado pequeños)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.5)  # Margen lateral restaurado
    section.right_margin = Inches(0.5)  # Margen lateral restaurado

def set_table_borders(table):
    tbl = table._element
    tblPr = tbl.xpath("w:tblPr")
    if not tblPr:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    else:
        tblPr = tblPr[0]
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

# Función para eliminar espacios de párrafo pero mantener legibilidad
def remove_paragraph_spacing(paragraph):
    if paragraph:
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        
        # Eliminar espaciado pero mantener un mínimo para legibilidad
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '20')  # Valor pequeño pero no cero
        spacing.set(qn('w:after'), '20')   # Valor pequeño pero no cero
        pPr.append(spacing)

# Función para ajustar espacios en celdas
def adjust_cell_spacing(cell):
    # Configurar márgenes de la celda
    tcPr = cell._element.tcPr
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        cell._element.append(tcPr)
    
    # Ajustar márgenes internos (pequeños pero no cero)
    tblCellMar = OxmlElement('w:tcMar')
    margins = {'top': '40', 'left': '60', 'bottom': '40', 'right': '60'}
    for margin, value in margins.items():
        mar = OxmlElement(f'w:{margin}')
        mar.set(qn('w:w'), value)
        mar.set(qn('w:type'), 'dxa')
        tblCellMar.append(mar)
    
    # Reemplazar márgenes existentes o agregar nuevos
    existing_mar = tcPr.find(qn('w:tcMar'))
    if existing_mar is not None:
        tcPr.remove(existing_mar)
    tcPr.append(tblCellMar)
    
    # Establecer alineación vertical superior
    tcVAlign = OxmlElement('w:vAlign')
    tcVAlign.set(qn('w:val'), 'top')
    tcPr.append(tcVAlign)
    
    # Ajustar espacios en todos los párrafos de la celda
    for paragraph in cell.paragraphs:
        remove_paragraph_spacing(paragraph)

# Iterar a través de cada fila en el DataFrame
for index, row in df.iterrows():
    # Obtener el número de cajas para este código
    num_cajas = int(row.get('CONTEO_CAJAS', 1))  # Valor predeterminado de 1 si no existe la columna
    
    # Para cada caja, crear una tabla
    for num_caja in range(1, num_cajas + 1):
        # Tabla principal (2 columnas, 2 filas)
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.columns[0].width = Inches(3)
        table.columns[1].width = Inches(3)
        set_table_borders(table)
        
        # Ajustar espacios en todas las celdas de la tabla principal
        for row_cells in table.rows:
            for cell in row_cells.cells:
                adjust_cell_spacing(cell)

        # --- Columna izquierda (0,0): imagen arriba, tabla de 6 celdas debajo ---
        left_cell = table.cell(0, 0)
        
        if 'IMAGEN' in row and pd.notna(row['IMAGEN']):
            try:
                # Agregar imagen con espaciado mínimo
                p = left_cell.paragraphs[0]  # Usar el primer párrafo existente
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p.add_run()
                run.add_picture(str(row['IMAGEN']), width=Inches(2.5))
            except Exception:
                p = left_cell.paragraphs[0]
                p.text = "Imagen no encontrada"
        
        # Tabla de 6 filas debajo de la imagen
        data_table = left_cell.add_table(rows=6, cols=1)
        data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(data_table)
        
        # Ajustar espacios en todas las celdas de la tabla de datos
        for row_cells in data_table.rows:
            for cell in row_cells.cells:
                adjust_cell_spacing(cell)
        
        data_table.cell(0, 0).text = f"CÓDIGO: {row['CODIGO']}"
        data_table.cell(1, 0).text = "REVISADO POR:"
        data_table.cell(2, 0).text = f"DESCRIPCIÓN: {row['DESCRIPCION']}"
        data_table.cell(3, 0).text = f"CAJA {num_caja} DE {num_cajas}"  # Aquí se actualiza el número de caja
        data_table.cell(4, 0).text = "RECEPCIONADO:"
        data_table.cell(5, 0).text = f"CANTIDAD: {row['CANTIDAD']}"

        # --- Columna derecha (0,1): tabla interna de 22 filas x 4 columnas ---
        right_cell = table.cell(0, 1)
        
        # Eliminar el párrafo existente en la celda derecha para quitar la línea
        if right_cell.paragraphs:
            p = right_cell.paragraphs[0]
            p._element.getparent().remove(p._element)
            p._p = None
        
        # Usar el primer párrafo existente para la tabla
        inner_table = right_cell.add_table(rows=22, cols=4)
        inner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(inner_table)
        
        # Ajustar espacios en todas las celdas de la tabla interna
        for row_cells in inner_table.rows:
            for cell in row_cells.cells:
                adjust_cell_spacing(cell)
        
        headers = ['FECHA DD/MM/AA', 'CANT.', 'RP.', 'FIRMA']
        for j, header in enumerate(headers):
            inner_table.cell(0, j).text = header
        
        for i in range(1, 22):
            for j in range(4):
                inner_table.cell(i, j).text = ""

        # Segunda fila: celda fusionada para "OBSERVACIÓN:" con altura aumentada
        obs_cell = table.cell(1, 0)
        obs_cell.merge(table.cell(1, 1))
        obs_cell.text = "OBSERVACIÓN:"
        
        # Aumentar la altura de la celda de observación (equivalente a 2 celdas)
        tr = obs_cell._element.getparent()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '800')  # Valor en twips (1/20 de punto) - aumentado para que sea más alto
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr = tr.get_or_add_trPr()
        trPr.append(trHeight)
        
        # Agregar salto de página después de cada tabla (excepto la última)
        if not (index == len(df) - 1 and num_caja == num_cajas):
            doc.add_page_break()

# Guardar el documento Word modificado
doc.save('recursos/documento.docx')